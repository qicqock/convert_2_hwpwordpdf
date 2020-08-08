from docx import *
from docx.shared import Pt
import os
from docx.oxml.ns import qn
from docx2pdf import convert

# 실행 예시:
# python txt_to_word.py --> 받아올 텍스트파일입력(txt)
# 이 파이썬 코드는 텍스트 형식의 내용을 받아 속기사가 작성하는 회의록 형식의 워드파일로 만들어주는 코드입니다.
# 텍스트를 워드로 변환하는 과정은 python-docx,워드를 pdf로 변환하는 과정은 docx2pdf 라는 파이썬 라이브러리를 사용했습니다.
# python-docx로 조작할수 없는 word파일의 설정이 있기때문에 이 부분은 standard.docx라는 기존의 파일을 이용하여 미리 설정을 해두었습니다.
# 주석은 해당하는 코드 밑에 작성했습니다.
# 순서
# 1: 다단을 설정해주는 함수 정의(set_number_of_columns)
# 2: 텍스트를 워드와 pdf로 변환하는 함수 정의(trans)
#   2.1: 텍스트 파일 불러오기 및 예외처리
#   2.2: 불러온 텍스트파일에 특수기호 처리
#   2.3: python-docx 사용하여 워드 파일 불러옴
#   2.4: 특수기호 처리한 텍스트를 워드파일에 추가
#   2.5: 글씨크기 및 폰트 변경
#   2.6: # 1)의 함수 이용하여 다단설정
#   2.7: 워드 docx 파일 저장
#   2.8: pdf 변환
# 3: trans 함수 실행

# 준비사항
# python 라이브러리 관리도구인 pip 설치
# pip를 통해 python-docx 와 docx2pdf 설치
# 명령어
# pip install python-docx
# pip install docx2pdf





WNS_COLS_NUM = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num"


# 1)
def set_number_of_columns(section, cols):
    section._sectPr.xpath("./w:cols")[0].set(WNS_COLS_NUM, str(cols))
    # 다단을 설정해주는 함수. (텍스트를 반으로 나눠줌)
    # 적용할 section과 다단의 개수를 파라미터로 받음.
    # 페이지 가운데를 기준으로 글을 나눠줌.
    # section의 기능을 사용


# 2)
def trans():
    # 2.1)
    print("Enter the txt file name")
    txt_name = input()

    try:
        f = open(txt_name, 'rt', encoding='utf-8')
        # 텍스트 파일을 열어서 확인
    except FileNotFoundError:
        print("Cannot find the file")
        return
        # 파일이 없어서 열수없을때 오류처리

    test = f.read()
    # 텍스트파일에 있는 내용을 test로 받아옴
    f.close()

    # 2.2)
    for i in range(0, len(test) - 2):
        if test[i] == test[i + 1] == '\n':
            templist = list(test)
            templist[i + 1] = "○"
            test = "".join(templist)
            # 문단이 끝나는곳을 표시하기 위해  두번역속 개행이 나오면 "^"을 넣어서 표시했음
            # 파이썬에서 string은 수정할수 없기떄문에 list로 변환후 그 list를 다시 스트링으로 변환하는 과정을 거침

    test = "○" + test
    # 처음 의원이름앞에 특수기호 추가

    # 2.3)
    document = Document('standard.docx')
    # python-docx 사용
    # 파일 형식을 맞추기 위해 standard.docx를 불러옴

    document._body.clear_content()
    # 원래 있던 standard파일에 내용을 없애줌

    # 2.4)
    alist = test.split('\n')
    #test라는 스트링을 개행단위를 기준으로 list형식으로 바꿈

    for i, v in enumerate(alist):
        try:
            if v[0] == "○":
                para = document.add_paragraph()
                para.paragraph_format.line_spacing = 1.2
                run = para.add_run(v)
                run.bold = True
                para.add_run("  " + alist[i + 1])
                del alist[i + 1]
                # 문장에서 첫글자가 '○'일때 그 문장을 굵은 글씨체 처리
                # 그 바로 뒤에 오는 문장은 개행하지 않음
                # 개행간격을 1.2로 설정
            else:
                para = document.add_paragraph()
                run = para.add_run("  " + v)
                para.paragraph_format.line_spacing = 1.2
                # 문장의 앞에 들여쓰기 "  "추가
                # 개행간격을 1.2로 설정
        except IndexError:
            continue

    # 2.5)
    style = document.styles['Normal']
    font = style.font
    font.name = '함초롱바탕'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '함초롱바탕')
    font.size = Pt(10)
    # 스타일을 이용하여 글씨크기 변경 및 폰트 변경

    # 2.6)
    document.add_page_break()
    section = document.sections[0]
    set_number_of_columns(section, 2)
    # 다단설정 함수를 사용하여 2개의 다단으로 만들어줌

    # 2.7)
    filename = os.path.splitext(txt_name)[0]
    docxfilename = filename + ".docx"
    document.save(docxfilename)
    # 원래 파일이름 그대로로 확장자만 바꿔서 저장

    # 2.8)
    convert(docxfilename)
    # pdf로 변환

# 3)
trans()
